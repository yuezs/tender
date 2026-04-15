from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(r"E:\tender\demo_materials\leader_demo_smart_park")


def ensure_dirs() -> None:
    for path in [
        ROOT / "knowledge" / "company_profile",
        ROOT / "knowledge" / "business_info",
        ROOT / "knowledge" / "qualifications",
        ROOT / "knowledge" / "templates",
        ROOT / "knowledge" / "project_cases",
        ROOT / "tender",
        ROOT / "proposal",
    ]:
        path.mkdir(parents=True, exist_ok=True)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].font.name = "Microsoft YaHei"


def build_doc(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.name = "Microsoft YaHei"
    run2.font.size = Pt(10.5)

    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(str(path))


def build_text_pdf(path: Path, lines: list[str]) -> None:
    objects: list[bytes] = []
    stream_lines = ["BT", "/F1 11 Tf"]
    y = 760
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream_lines.append(f"72 {y} Td ({safe}) Tj")
        stream_lines.append("0 -18 Td")
        y -= 18
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objects.append(
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>\nendobj\n"
    )
    objects.append(
        f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("latin-1")
        + stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)

    startxref = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    pdf.extend(
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{startxref}\n%%EOF\n".encode(
            "latin-1"
        )
    )
    path.write_bytes(pdf)


def main() -> None:
    ensure_dirs()

    build_doc(
        ROOT / "knowledge" / "company_profile" / "company_profile_demo.docx",
        "星城数智科技有限公司公司简介",
        "演示知识库资料 / category: company_profile",
        [
            (
                "公司概况",
                [
                    "星城数智科技有限公司成立于 2016 年，专注智慧园区、园区安防、能耗监测、物联感知与运营平台建设。",
                    "公司总部位于长沙，在武汉、合肥、西安设有交付支持团队，现有员工 168 人，其中研发与实施人员 112 人。",
                    "公司累计交付 40 余个园区数字化项目，具备平台研发、项目实施、系统集成和运维服务一体化能力。",
                ],
            ),
            (
                "核心能力",
                [
                    "具备园区一张图、人员车辆管理、视频监控接入、能耗监测、移动巡检和第三方系统集成等成熟模块。",
                    "支持私有化部署、统一身份认证、日志审计和标准 API 集成。",
                    "项目实施采用需求调研、蓝图设计、分阶段上线、试运行、验收的标准交付路径。",
                ],
            ),
            (
                "项目适配优势",
                [
                    "已有两个以上园区平台类成功案例，可直接支撑本次招标文件的类似项目业绩要求。",
                    "具备软件类项目常用质量与信息安全体系资质，可支撑商务评分和资格响应。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "business_info" / "business_info_demo.docx",
        "星城数智科技有限公司商务信息与售后承诺",
        "演示知识库资料 / category: business_info",
        [
            (
                "商务响应要点",
                [
                    "接受公开招标方式和 6 个月总工期要求。",
                    "提供项目经理、技术负责人、实施工程师、测试工程师和运维工程师组成的标准交付团队。",
                    "按要求提交技术方案、商务响应文件、报价文件和资格证明文件。",
                ],
            ),
            (
                "售后服务",
                [
                    "提供不少于 12 个月质保服务，重大故障 2 小时响应，24 小时内到场。",
                    "提供管理员、业务用户和运维人员三类培训材料与现场培训支持。",
                    "提供月度巡检、季度优化建议和年度系统健康检查。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "qualifications" / "qualifications_demo.docx",
        "星城数智科技有限公司资质证书与合规说明",
        "演示知识库资料 / category: qualifications",
        [
            (
                "基础资质",
                [
                    "企业为独立法人，经营范围覆盖软件开发、信息系统集成服务和智能化项目实施。",
                    "具备电子与智能化工程专业承包相关实施能力。",
                    "具备园区视频、门禁和安防接入项目经验。",
                ],
            ),
            (
                "体系认证",
                [
                    "已通过 ISO9001 质量管理体系认证。",
                    "已通过 ISO27001 信息安全管理体系认证。",
                    "已通过 ISO20000 信息技术服务管理体系认证。",
                    "已取得 CMMI 三级能力认证。",
                ],
            ),
            (
                "合规声明",
                [
                    "近三年无重大违法记录，无失信被执行情况，无重大质量安全责任事故。",
                    "可按项目要求提供营业执照、资质证书、认证证书和业绩材料。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "templates" / "templates_demo.docx",
        "智慧园区综合管理平台技术方案模板",
        "演示知识库资料 / category: templates",
        [
            (
                "总体设计模板",
                [
                    "采用应用层、服务层、数据层、设备接入层分层架构，支持私有化部署和横向扩展。",
                    "统一管理人员、车辆、设备、视频、能耗和事件数据。",
                    "具备统一认证、权限分级、日志审计、备份与恢复机制。",
                ],
            ),
            (
                "功能模块模板",
                [
                    "园区一张图、人员车辆管理、视频监控接入、能耗监测、移动巡检、系统集成。",
                    "支持与 OA、门禁、视频平台、能耗系统和单点登录平台对接。",
                ],
            ),
            (
                "实施计划模板",
                [
                    "第 1 月：需求调研、现状摸排、蓝图设计。",
                    "第 2-4 月：功能开发、设备接入联调、分模块测试。",
                    "第 5 月：试运行、培训、优化整改。",
                    "第 6 月：验收准备、资料归档、终验汇报。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "project_cases" / "project_cases_demo.docx",
        "智慧园区与园区安防项目案例集",
        "演示知识库资料 / category: project_cases",
        [
            (
                "案例一",
                [
                    "嘉衡高新区智慧园区综合管理平台项目，合同金额 568 万元，建设周期 5.5 个月。",
                    "完成 680 路视频接入、42 个门禁通道联动、260 个能耗计量点监测。",
                ],
            ),
            (
                "案例二",
                [
                    "临港智能制造园区综合治理平台项目，合同金额 632 万元，建设周期 6 个月。",
                    "实现 9000+ 人员档案统一管理和 3500+ 车辆通行规则管理。",
                ],
            ),
            (
                "可复用经验",
                [
                    "园区平台项目的关键成功点是前期数据标准梳理、接口范围界定和分阶段上线。",
                    "对视频监控接入、能耗监测和人员车辆管理组合需求，统一平台底座更利于交付和运维。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "tender" / "tender_demo.docx",
        "智慧园区综合管理平台建设项目示例招标文件",
        "演示上传文件 / 建议用于 tender 主链路",
        [
            (
                "项目基本信息",
                [
                    "项目名称：智慧园区综合管理平台建设项目。",
                    "项目编号：XC-GXQ-2026-01。",
                    "招标单位：星城高新技术产业园区管理委员会。",
                    "预算金额：人民币 480 万元。",
                    "项目工期：合同签订后 6 个月内完成建设并验收。",
                ],
            ),
            (
                "建设内容",
                [
                    "园区基础数据管理与园区一张图。",
                    "人员与车辆综合管理。",
                    "视频监控接入与告警联动。",
                    "能耗监测与分析。",
                    "移动端巡检与事件上报。",
                    "统一消息通知与第三方系统集成。",
                ],
            ),
            (
                "技术要求与资格要求",
                [
                    "采用微服务或模块化架构，支持私有化部署。",
                    "支持不少于 10000 人员档案、5000 车辆档案、800 路视频接入和 200 个计量点监测。",
                    "投标人须为独立法人，具备两个及以上类似项目业绩。",
                    "具备 ISO9001、ISO27001、ISO20000 或同类认证，CMMI 优先。",
                    "近三年无重大违法记录，不接受联合体投标。",
                ],
            ),
            (
                "评分与时间安排",
                [
                    "技术方案 60 分，商务部分 20 分，价格部分 20 分。",
                    "投标截止时间：2026 年 5 月 10 日 17:00。",
                    "开标时间：2026 年 5 月 11 日 09:30。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "proposal" / "proposal_demo.docx",
        "智慧园区综合管理平台建设项目标书演示稿",
        "基于演示知识库整理的参考标书稿",
        [
            (
                "项目理解与投标结论",
                [
                    "本项目与公司既有智慧园区平台产品、园区安防集成经验和能耗监测实施经验高度匹配，建议积极参与投标。",
                    "公司已具备独立法人资格、体系认证和园区平台类成功案例，可形成完整响应。",
                ],
            ),
            (
                "总体技术方案",
                [
                    "采用应用层、服务层、数据层、设备接入层分层架构，在私有化环境部署统一园区综合管理平台。",
                    "平台核心模块包括园区一张图、人员车辆管理、视频监控接入、能耗监测、移动巡检和系统集成。",
                ],
            ),
            (
                "实施计划与服务承诺",
                [
                    "按 6 个月工期分阶段交付，覆盖调研、开发、联调、试运行和验收。",
                    "提供不少于 12 个月质保服务，重大故障 2 小时响应，24 小时内到场。",
                ],
            ),
        ],
    )

    build_text_pdf(
        ROOT / "tender" / "tender_demo_text_pdf.pdf",
        [
            "Tender demo smart park platform project",
            "Project XC-GXQ-2026-01 owner Xingcheng Hi-tech Industrial Park Committee",
            "Budget 4.8 million RMB duration 6 months open tender",
            "Scope smart park map personnel vehicle video energy mobile app integration",
            "Requirements private deployment api integration mysql audit log high availability",
            "Qualification independent legal entity two similar projects iso9001 iso27001 iso20000 cmmi preferred",
            "Scoring technical 60 business 20 price 20",
            "Deadline 2026-05-10 17:00 opening 2026-05-11 09:30",
        ],
    )

    (ROOT / "README.md").write_text(
        "# 领导演示资料包\n\n"
        "## 目录\n"
        "- `knowledge/`：企业知识库示例资料\n"
        "- `tender/`：可上传的示例招标文件\n"
        "- `proposal/`：对应该项目的参考标书演示稿\n\n"
        "## 推荐演示顺序\n"
        "1. 在资料中心依次上传 `knowledge/` 下 5 份知识库文档\n"
        "2. 对 5 份文档执行处理\n"
        "3. 在招标处理页上传 `tender/tender_demo.docx`\n"
        "4. 执行解析、抽取、判断、生成主链路\n"
        "5. 在结果页展示提取字段、投标建议、目录和章节生成\n"
        "6. 对照 `proposal/proposal_demo.docx` 讲解预期输出风格\n\n"
        "## 说明\n"
        "- 这套资料主题统一为“智慧园区综合管理平台建设项目”，便于整链路演示\n"
        "- `tender/` 目录同时附带一个文本型 PDF 示例，可用于演示文本 PDF 解析能力\n",
        encoding="utf-8",
    )

    (ROOT / "DEMO_SCRIPT.md").write_text(
        "# 智慧园区项目演示清单\n\n"
        "## 推荐顺序\n"
        "1. 先在 `/knowledge` 上传并处理 5 份企业知识库资料\n"
        "2. 再在 `/tender` 上传 `tender_demo.docx` 跑主链路\n"
        "3. 最后到 `/results` 展示字段、结论、风险、目录和章节内容\n\n"
        "## 建议讲法\n"
        "- 先讲企业知识如何沉淀成可调用的知识库\n"
        "- 再讲招标文件如何被系统解析和判断\n"
        "- 最后讲标书初稿如何结合企业资料自动生成\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
