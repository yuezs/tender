from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(r"E:\tender\demo_materials\leader_demo_smart_water")


def ensure_dirs() -> None:
    for path in [
        ROOT / "knowledge" / "company_profile",
        ROOT / "knowledge" / "business_info",
        ROOT / "knowledge" / "qualifications",
        ROOT / "knowledge" / "templates",
        ROOT / "knowledge" / "project_cases",
        ROOT / "tender",
        ROOT / "proposal",
    ]:
        path.mkdir(parents=True, exist_ok=True)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].font.name = "Microsoft YaHei"


def build_doc(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    style_doc(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run.font.size = Pt(16)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run.font.size = Pt(10.5)

    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(str(path))


def build_text_pdf(path: Path, lines: list[str]) -> None:
    objects: list[bytes] = []
    stream_lines = ["BT", "/F1 11 Tf"]
    y = 760
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream_lines.append(f"72 {y} Td ({safe}) Tj")
        stream_lines.append("0 -18 Td")
        y -= 18
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objects.append(
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>\nendobj\n"
    )
    objects.append(
        f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("latin-1")
        + stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)

    startxref = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    pdf.extend(
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{startxref}\n%%EOF\n".encode(
            "latin-1"
        )
    )
    path.write_bytes(pdf)


KNOWLEDGE_UPLOAD_ITEMS = [
    {
        "title": "华澜数智科技有限公司公司简介",
        "category": "company_profile",
        "tags": "智慧水务,物联网平台,数据治理,县域项目",
        "industry": "水务,市政,政务",
        "path": "knowledge/company_profile/company_profile_demo.docx",
    },
    {
        "title": "华澜数智科技有限公司商务信息与服务承诺",
        "category": "business_info",
        "tags": "商务响应,售后服务,交付保障,培训运维",
        "industry": "水务,市政,公共服务",
        "path": "knowledge/business_info/business_info_demo.docx",
    },
    {
        "title": "华澜数智科技有限公司资质证书与合规说明",
        "category": "qualifications",
        "tags": "ISO9001,信息安全,系统集成,合规资质",
        "industry": "水务,信息化,市政",
        "path": "knowledge/qualifications/qualifications_demo.docx",
    },
    {
        "title": "智慧水务运行监测与调度平台技术方案模板",
        "category": "templates",
        "tags": "技术方案,平台架构,实施计划,数据中台",
        "industry": "水务,数字政府,市政",
        "path": "knowledge/templates/templates_demo.docx",
    },
    {
        "title": "智慧水务项目案例集",
        "category": "project_cases",
        "tags": "项目案例,泵站监测,管网监测,漏损治理",
        "industry": "水务,环保,市政",
        "path": "knowledge/project_cases/project_cases_demo.docx",
    },
]


def write_metadata_files() -> None:
    csv_path = ROOT / "UPLOAD_METADATA.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["title", "category", "tags", "industry", "path"])
        writer.writeheader()
        for item in KNOWLEDGE_UPLOAD_ITEMS:
            writer.writerow(item)

    lines = [
        "# 演示资料上传清单",
        "",
        "## 知识库资料",
        "",
        "| 标题 | 分类 | 标签 | 行业 | 文件路径 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for item in KNOWLEDGE_UPLOAD_ITEMS:
        lines.append(
            f"| {item['title']} | {item['category']} | {item['tags']} | {item['industry']} | `{item['path']}` |"
        )

    lines.extend(
        [
            "",
            "## 招标文件",
            "",
            "- 标题：县域智慧水务运行监测与调度平台建设项目示例招标文件",
            "- 推荐上传路径：`tender/tender_demo.docx`",
            "- 备用文本型 PDF：`tender/tender_demo_text_pdf.pdf`",
            "- 推荐标签：智慧水务,运行调度,监测平台,公开招标",
            "- 推荐行业：水务,市政,政务",
            "",
            "## 标书演示稿",
            "",
            "- 文件：`proposal/proposal_demo.docx`",
            "- 用途：结果页演示、章节生成讲解、全文导出对照",
        ]
    )
    (ROOT / "UPLOAD_METADATA.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()

    build_doc(
        ROOT / "knowledge" / "company_profile" / "company_profile_demo.docx",
        "华澜数智科技有限公司公司简介",
        "演示知识库资料 / category: company_profile",
        [
            (
                "公司概况",
                [
                    "华澜数智科技有限公司成立于 2017 年，专注智慧水务、物联网感知、运行监测与数据治理平台建设。",
                    "公司总部位于南京，在苏州、合肥、武汉设有交付支持团队，现有员工 152 人，其中研发和实施人员 98 人。",
                    "公司累计服务 30 余个市县级供排水信息化项目，具备平台研发、项目交付、系统集成和运维服务一体化能力。",
                ],
            ),
            (
                "核心能力",
                [
                    "具备水厂监测、泵站监测、供水管网监测、漏损分析、告警处置、调度驾驶舱和移动巡检等成熟模块。",
                    "支持私有化部署、统一身份认证、审计日志、地图展示和第三方系统标准接口集成。",
                    "在县域项目中具备从需求调研、蓝图设计、分阶段建设到试运行和验收的完整实施经验。",
                ],
            ),
            (
                "与本项目的适配优势",
                [
                    "具备水务运行监测与调度平台类近似案例，可直接支撑类似业绩要求。",
                    "既有平台能力覆盖招标文件中的监测、预警、分析、调度和报表等主要业务场景。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "business_info" / "business_info_demo.docx",
        "华澜数智科技有限公司商务信息与服务承诺",
        "演示知识库资料 / category: business_info",
        [
            (
                "商务响应要点",
                [
                    "接受公开招标方式和 8 个月总工期要求。",
                    "提供项目经理、架构师、实施工程师、测试工程师和运维工程师组成的标准交付团队。",
                    "支持按招标要求提交技术响应文件、商务响应文件、报价文件和资格证明材料。",
                ],
            ),
            (
                "服务承诺",
                [
                    "提供不少于 12 个月质保服务，重大故障 2 小时响应，24 小时内到场。",
                    "提供管理员、调度人员和运维人员三类培训材料及现场培训支持。",
                    "提供月度巡检、季度优化建议和年度系统健康检查服务。",
                ],
            ),
            (
                "交付保障",
                [
                    "采用里程碑交付和周报机制，保障项目进度与风险可控。",
                    "提供数据迁移、接口联调和试运行支持，确保系统平滑上线。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "qualifications" / "qualifications_demo.docx",
        "华澜数智科技有限公司资质证书与合规说明",
        "演示知识库资料 / category: qualifications",
        [
            (
                "基础资质",
                [
                    "企业为独立法人，经营范围覆盖软件开发、信息系统集成服务和智慧水务相关项目实施。",
                    "具备电子与智能化系统集成相关实施能力。",
                    "具备供排水监测、泵站监测和行业平台项目建设经验。",
                ],
            ),
            (
                "体系认证",
                [
                    "已通过 ISO9001 质量管理体系认证。",
                    "已通过 ISO27001 信息安全管理体系认证。",
                    "已通过 ISO20000 信息技术服务管理体系认证。",
                    "已取得 CMMI 三级能力认证。",
                ],
            ),
            (
                "合规声明",
                [
                    "近三年无重大违法记录，无失信被执行情况，无重大质量安全责任事故。",
                    "可按项目要求提供营业执照、资质证书、认证证书和项目业绩材料。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "templates" / "templates_demo.docx",
        "智慧水务运行监测与调度平台技术方案模板",
        "演示知识库资料 / category: templates",
        [
            (
                "总体设计模板",
                [
                    "采用感知接入层、数据治理层、业务服务层、驾驶舱展示层四层架构，支持私有化部署和分阶段扩展。",
                    "统一管理水厂、泵站、管网、告警、工单和调度数据，形成一体化运行监测平台。",
                    "具备统一认证、权限分级、日志审计、备份恢复和高可用部署方案。",
                ],
            ),
            (
                "功能模块模板",
                [
                    "监测总览、泵站监测、水厂监测、管网监测、漏损分析、告警处置、调度指挥、报表统计、移动巡检。",
                    "支持与 SCADA、GIS、短信平台、视频监控系统和第三方业务系统接口对接。",
                ],
            ),
            (
                "实施计划模板",
                [
                    "第 1 月：需求调研、现状摸排、指标梳理、蓝图设计。",
                    "第 2-5 月：平台开发、设备接入、数据治理、接口联调和分模块测试。",
                    "第 6-7 月：试运行、培训、优化整改、数据校核。",
                    "第 8 月：验收准备、资料归档、终验汇报。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "knowledge" / "project_cases" / "project_cases_demo.docx",
        "智慧水务项目案例集",
        "演示知识库资料 / category: project_cases",
        [
            (
                "案例一",
                [
                    "江北新区智慧供水运行监测平台项目，合同金额 486 万元，建设周期 7 个月。",
                    "完成 12 座泵站、3 座水厂、180 个管网监测点和 420 路视频监控数据接入。",
                ],
            ),
            (
                "案例二",
                [
                    "临河县智慧水务综合调度平台项目，合同金额 528 万元，建设周期 8 个月。",
                    "实现管网监测、漏损分析、告警闭环和调度驾驶舱联动，覆盖 16 个乡镇供水区域。",
                ],
            ),
            (
                "可复用经验",
                [
                    "县域水务项目成功关键在于前期数据标准梳理、监测点位核查和接口边界定义。",
                    "对监测、预警、调度、报表组合需求，统一平台底座更利于交付、运维和后续扩展。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "tender" / "tender_demo.docx",
        "县域智慧水务运行监测与调度平台建设项目示例招标文件",
        "演示上传文件 / 建议用于 tender 主链路",
        [
            (
                "项目基本信息",
                [
                    "项目名称：县域智慧水务运行监测与调度平台建设项目。",
                    "项目编号：HL-SW-2026-03。",
                    "招标人：清源县水务局。",
                    "预算金额：人民币 518 万元。",
                    "建设周期：8 个月。",
                ],
            ),
            (
                "建设范围",
                [
                    "建设监测总览、泵站监测、水厂监测、管网监测、漏损分析、告警处置、调度指挥和报表统计模块。",
                    "接入泵站、水厂、管网监测点、视频监控和短信告警等数据源。",
                    "完成平台部署、数据治理、接口联调、培训、试运行和验收工作。",
                ],
            ),
            (
                "主要技术要求",
                [
                    "系统支持私有化部署，采用 B/S 架构，数据库使用 MySQL。",
                    "支持统一身份认证、操作日志审计、告警规则配置、权限分级和高可用部署。",
                    "支持与现有 GIS、SCADA、短信平台和视频监控系统对接。",
                ],
            ),
            (
                "资格与评分要求",
                [
                    "投标人须为独立法人，具备软件研发和系统集成能力。",
                    "近三年具有不少于 2 个智慧水务或行业监测平台类项目业绩。",
                    "评分权重：技术 60 分，商务 20 分，价格 20 分。",
                ],
            ),
        ],
    )

    build_doc(
        ROOT / "proposal" / "proposal_demo.docx",
        "县域智慧水务运行监测与调度平台建设项目标书演示稿",
        "基于演示知识库整理的参考标书稿",
        [
            (
                "项目理解与投标结论",
                [
                    "本项目与公司既有智慧水务监测平台产品、泵站监测交付经验和县域项目实施能力高度匹配，建议积极参与投标。",
                    "公司已具备独立法人资格、体系认证和水务平台类成功案例，可形成完整资格与技术响应。",
                ],
            ),
            (
                "总体技术方案",
                [
                    "项目拟采用感知接入、数据治理、业务服务、驾驶舱展示四层架构，实现监测、预警、分析和调度一体化。",
                    "通过统一平台整合水厂、泵站、管网、告警和视频等数据，形成县域水务运行监测与调度中心。",
                ],
            ),
            (
                "实施与服务承诺",
                [
                    "采用 8 个月分阶段实施路径，覆盖调研、建设、联调、试运行和验收全过程。",
                    "提供 12 个月质保、现场培训、月度巡检和持续优化建议，保障系统稳定运行。",
                ],
            ),
        ],
    )

    build_text_pdf(
        ROOT / "tender" / "tender_demo_text_pdf.pdf",
        [
            "Smart water monitoring and dispatch platform tender demo",
            "Project HL-SW-2026-03 owner Qingyuan County Water Affairs Bureau",
            "Budget 5.18 million RMB duration 8 months open tender",
            "Scope plant monitoring pump station monitoring pipe network alerts dispatch reports",
            "Requirements private deployment mysql api integration audit log high availability",
            "Qualification two similar projects iso9001 iso27001 iso20000 cmmi preferred",
            "Scoring technical 60 business 20 price 20",
            "Deadline 2026-06-18 17:00 opening 2026-06-19 09:30",
        ],
    )

    write_metadata_files()

    (ROOT / "README.md").write_text(
        "\n".join(
            [
                "# 领导演示资料包（智慧水务版）",
                "",
                "## 目录",
                "- `knowledge/`：企业知识库示例资料",
                "- `tender/`：可上传的示例招标文件",
                "- `proposal/`：对应该项目的参考标书演示稿",
                "- `UPLOAD_METADATA.md` / `UPLOAD_METADATA.csv`：上传时可直接填写的标题、分类、标签、行业清单",
                "",
                "## 推荐演示顺序",
                "1. 在资料中心依次上传 `knowledge/` 下 5 份知识库文档",
                "2. 按 `UPLOAD_METADATA.*` 中的标题、分类、标签、行业录入资料信息",
                "3. 对 5 份文档执行处理",
                "4. 在招标处理页上传 `tender/tender_demo.docx`",
                "5. 执行解析、抽取、判断、生成主链路",
                "6. 在结果页展示提取字段、投标建议、目录和章节生成",
                "7. 对照 `proposal/proposal_demo.docx` 讲解预期输出风格",
                "",
                "## 说明",
                "- 这套资料主题统一为“县域智慧水务运行监测与调度平台建设项目”，适合政企类场景演示。",
                "- `tender/` 目录同时附带一个文本型 PDF 示例，可用于演示文本 PDF 解析能力。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    (ROOT / "DEMO_SCRIPT.md").write_text(
        "\n".join(
            [
                "# 智慧水务项目演示清单",
                "",
                "## 一、演示目标",
                "",
                "用一套统一主题的资料，完整展示：",
                "",
                "1. 企业知识库上传与处理",
                "2. 招标文件上传与解析",
                "3. 核心字段抽取",
                "4. 是否建议投标判断",
                "5. 标书初稿生成与结果审阅",
                "",
                "## 二、资料对应关系",
                "",
                "### 1. 企业知识库",
                "",
                "- `knowledge/company_profile/company_profile_demo.docx`",
                "  - 用于展示企业概况、核心能力与项目适配性",
                "- `knowledge/business_info/business_info_demo.docx`",
                "  - 用于展示商务响应、交付与售后服务承诺",
                "- `knowledge/qualifications/qualifications_demo.docx`",
                "  - 用于展示资格证明、体系认证与合规说明",
                "- `knowledge/templates/templates_demo.docx`",
                "  - 用于展示技术方案、架构和实施计划模板",
                "- `knowledge/project_cases/project_cases_demo.docx`",
                "  - 用于展示类似项目案例和可复用经验",
                "",
                "### 2. 招标文件",
                "",
                "- `tender/tender_demo.docx`",
                "  - 建议用于主链路上传、解析、抽取和判断",
                "- `tender/tender_demo_text_pdf.pdf`",
                "  - 用于演示文本型 PDF 可解析，扫描件暂不支持",
                "",
                "### 3. 标书演示稿",
                "",
                "- `proposal/proposal_demo.docx`",
                "  - 用于对照讲解生成结果结构、章节组织和表达风格",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
