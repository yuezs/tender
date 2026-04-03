import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document

from core.config import settings


class ProposalDocumentService:
    OUTLINE_PLACEHOLDER = "待根据目录逐节生成正文。"
    FULLTEXT_PLACEHOLDER = "当前小节尚未生成正文，请后续补充。"

    def __init__(self, storage_root: Path | None = None) -> None:
        self.storage_root = Path(storage_root or settings.storage_root)
        self.documents_dir = self.storage_root / "tender" / "documents"
        self.documents_dir.mkdir(parents=True, exist_ok=True)

    def export(self, *, file_id: str, tender_record: dict, generate_result: dict) -> dict:
        document_id = uuid4().hex
        target_dir = self.documents_dir / file_id
        target_dir.mkdir(parents=True, exist_ok=True)

        extract_result = tender_record.get("extract_result", {}) or {}
        project_name = self._normalize_text(extract_result.get("project_name"), "标书目录初稿")
        tender_company = self._normalize_text(extract_result.get("tender_company"), "待确认招标单位")
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        file_name = f"{self._sanitize_file_name(project_name)}-{document_id[:8]}.docx"
        storage_path = target_dir / file_name

        proposal_outline = self._normalize_outline(generate_result)
        section_contents = self._normalize_section_contents(generate_result)

        document = Document()
        document.add_heading("标书目录初稿", level=0)
        document.add_paragraph(f"项目名称：{project_name}")
        document.add_paragraph(f"招标单位：{tender_company}")
        document.add_paragraph(f"生成时间：{generated_at}")
        document.add_page_break()

        document.add_heading("标书目录", level=1)
        self._append_outline_overview(document, proposal_outline)
        document.add_page_break()
        self._append_outline_sections(document, proposal_outline, section_contents)

        document.save(storage_path)

        return {
            "document_id": document_id,
            "file_name": file_name,
            "storage_path": str(storage_path),
            "download_url": f"/api/tender/documents/{document_id}/download",
            "generated_at": generated_at,
        }

    def export_full_text(self, *, file_id: str, tender_record: dict, generate_result: dict) -> dict:
        document_id = uuid4().hex
        target_dir = self.documents_dir / file_id
        target_dir.mkdir(parents=True, exist_ok=True)

        extract_result = tender_record.get("extract_result", {}) or {}
        project_name = self._normalize_text(extract_result.get("project_name"), "标书正文草稿")
        tender_company = self._normalize_text(extract_result.get("tender_company"), "待确认招标单位")
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        file_name = f"{self._sanitize_file_name(project_name)}-全文-{document_id[:8]}.docx"
        storage_path = target_dir / file_name

        proposal_outline = self._normalize_outline(generate_result)
        section_contents = self._normalize_section_contents(generate_result)

        document = Document()
        document.add_heading("标书正文草稿", level=0)
        document.add_paragraph(f"项目名称：{project_name}")
        document.add_paragraph(f"招标单位：{tender_company}")
        document.add_paragraph(f"生成时间：{generated_at}")
        document.add_page_break()
        self._append_full_text_sections(document, proposal_outline, section_contents)
        document.save(storage_path)

        return {
            "document_id": document_id,
            "file_name": file_name,
            "storage_path": str(storage_path),
            "download_url": f"/api/tender/documents/{document_id}/download",
            "generated_at": generated_at,
        }

    def _normalize_outline(self, generate_result: dict) -> list[dict]:
        raw_outline = generate_result.get("proposal_outline")
        if not isinstance(raw_outline, list):
            return []

        normalized: list[dict] = []
        for index, item in enumerate(raw_outline, start=1):
            if not isinstance(item, dict):
                continue

            title = self._normalize_text(item.get("title"), "")
            if not title:
                continue

            section_id = self._normalize_text(item.get("section_id"), str(index))
            purpose = self._normalize_text(item.get("purpose"), "待补充章节目标说明。")
            children: list[dict] = []

            for child_index, child in enumerate(item.get("children", []), start=1):
                if not isinstance(child, dict):
                    continue
                child_title = self._normalize_text(child.get("title"), "")
                if not child_title:
                    continue
                child_id = self._normalize_text(child.get("section_id"), f"{section_id}.{child_index}")
                children.append(
                    {
                        "section_id": child_id,
                        "title": child_title,
                        "purpose": self._normalize_text(child.get("purpose"), "待补充小节目标说明。"),
                        "writing_points": self._normalize_points(child.get("writing_points")),
                    }
                )

            normalized.append(
                {
                    "section_id": section_id,
                    "title": title,
                    "purpose": purpose,
                    "children": children,
                }
            )

        return normalized

    def _normalize_section_contents(self, generate_result: dict) -> dict[str, dict]:
        raw_contents = generate_result.get("section_contents")
        if not isinstance(raw_contents, dict):
            return {}

        normalized: dict[str, dict] = {}
        for section_id, item in raw_contents.items():
            if not isinstance(item, dict):
                continue
            normalized[str(section_id)] = {
                "section_id": self._normalize_text(item.get("section_id"), str(section_id)),
                "parent_section_id": self._normalize_text(item.get("parent_section_id"), ""),
                "title": self._normalize_text(item.get("title"), "未命名小节"),
                "status": self._normalize_text(item.get("status"), "pending"),
                "content": self._normalize_text(item.get("content"), ""),
                "error_message": self._normalize_text(item.get("error_message"), ""),
                "updated_at": self._normalize_text(item.get("updated_at"), ""),
            }
        return normalized

    def _append_outline_overview(self, document: Document, proposal_outline: list[dict]) -> None:
        for chapter in proposal_outline:
            chapter_prefix = self._prefix(chapter.get("section_id"))
            chapter_title = self._normalize_text(chapter.get("title"), "未命名章节")
            chapter_purpose = self._normalize_text(chapter.get("purpose"), "")

            document.add_paragraph(f"{chapter_prefix}{chapter_title}")
            if chapter_purpose:
                document.add_paragraph(f"章节目标：{chapter_purpose}")

            for child in chapter.get("children", []):
                child_prefix = self._prefix(child.get("section_id"))
                child_title = self._normalize_text(child.get("title"), "未命名小节")
                document.add_paragraph(f"{child_prefix}{child_title}")
                for point in child.get("writing_points", [])[:3]:
                    document.add_paragraph(str(point), style="List Bullet")
            document.add_paragraph("")

    def _append_outline_sections(
        self,
        document: Document,
        proposal_outline: list[dict],
        section_contents: dict[str, dict],
    ) -> None:
        for chapter in proposal_outline:
            chapter_id = self._normalize_text(chapter.get("section_id"), "")
            chapter_prefix = self._prefix(chapter_id)
            chapter_title = self._normalize_text(chapter.get("title"), "未命名章节")
            chapter_purpose = self._normalize_text(chapter.get("purpose"), "")

            document.add_heading(f"{chapter_prefix}{chapter_title}", level=1)
            if chapter_purpose:
                document.add_paragraph(f"章节目标：{chapter_purpose}")

            for child in chapter.get("children", []):
                child_id = self._normalize_text(child.get("section_id"), "")
                child_prefix = self._prefix(child_id)
                child_title = self._normalize_text(child.get("title"), "未命名小节")
                child_purpose = self._normalize_text(child.get("purpose"), "")
                child_content = section_contents.get(child_id, {})
                body = self._normalize_text(child_content.get("content"), self.OUTLINE_PLACEHOLDER)

                document.add_heading(f"{child_prefix}{child_title}", level=2)
                if child_purpose:
                    document.add_paragraph(f"小节目标：{child_purpose}")
                document.add_paragraph(body)
                document.add_paragraph("写作要点：")
                for point in child.get("writing_points", [])[:5]:
                    document.add_paragraph(str(point), style="List Bullet")

            document.add_paragraph("")

    def _append_full_text_sections(
        self,
        document: Document,
        proposal_outline: list[dict],
        section_contents: dict[str, dict],
    ) -> None:
        for chapter in proposal_outline:
            chapter_id = self._normalize_text(chapter.get("section_id"), "")
            chapter_prefix = self._prefix(chapter_id)
            chapter_title = self._normalize_text(chapter.get("title"), "未命名章节")
            chapter_purpose = self._normalize_text(chapter.get("purpose"), "")

            document.add_heading(f"{chapter_prefix}{chapter_title}", level=1)
            if chapter_purpose:
                document.add_paragraph(f"章节目标：{chapter_purpose}")

            for child in chapter.get("children", []):
                child_id = self._normalize_text(child.get("section_id"), "")
                child_prefix = self._prefix(child_id)
                child_title = self._normalize_text(child.get("title"), "未命名小节")
                child_purpose = self._normalize_text(child.get("purpose"), "")
                child_content = section_contents.get(child_id, {})
                body = self._build_full_text_body(child_content)

                document.add_heading(f"{child_prefix}{child_title}", level=2)
                if child_purpose:
                    document.add_paragraph(f"小节目标：{child_purpose}")
                document.add_paragraph(body)

            document.add_paragraph("")

    def _build_full_text_body(self, child_content: dict) -> str:
        status = self._normalize_text(child_content.get("status"), "pending")
        content = self._normalize_text(child_content.get("content"), "")
        error_message = self._normalize_text(child_content.get("error_message"), "")

        if content:
            return content
        if status == "error" and error_message:
            return f"当前小节生成失败：{error_message}"
        return self.FULLTEXT_PLACEHOLDER

    def _normalize_points(self, value: object) -> list[str]:
        if not isinstance(value, list):
            return [self.OUTLINE_PLACEHOLDER]

        points = [str(item).strip() for item in value if isinstance(item, str) and item.strip()][:5]
        return points or [self.OUTLINE_PLACEHOLDER]

    def _prefix(self, section_id: object) -> str:
        normalized = self._normalize_text(section_id, "")
        return f"{normalized} " if normalized else ""

    def _sanitize_file_name(self, value: str) -> str:
        cleaned = re.sub(r'[\\/:*?"<>|]+', "-", value).strip()
        cleaned = re.sub(r"\s+", "-", cleaned)
        return cleaned[:80] or "proposal"

    def _normalize_text(self, value: object, fallback: str) -> str:
        if isinstance(value, str):
            cleaned = value.strip()
            return cleaned or fallback
        return fallback
